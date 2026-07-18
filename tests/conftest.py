"""Cấu hình test: chế độ local, thư mục dữ liệu tạm, bộ chấm mock."""
from __future__ import annotations

import io
import os
import tempfile

_TMP = tempfile.mkdtemp(prefix="dnu-test-")
os.environ.update({
    "APP_MODE": "local",
    "DATA_DIR": _TMP,
    "GRADER": "mock",
    "SECRET_KEY": "test-secret",
    "CRON_TOKEN": "test-cron",
    # Hạn nộp đặt xa trong tương lai để test không phụ thuộc đồng hồ hệ thống
    # (test nào cần "quá hạn" tự đặt lại deadline trong DB).
    "DEADLINE": "2099-12-31T17:00:00+07:00",
    "OPEN_AT": "2020-01-01T00:00:00+07:00",
})

import pytest  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from app.config import reset_settings  # noqa: E402

reset_settings()

from app.main import create_app  # noqa: E402

_app = create_app()

USERS = [
    {"id": "u-gv1", "ma_gv": "GV001", "ho_ten": "Nguyễn Văn An", "email": "gv001@dainam.edu.vn",
     "khoa": "CNTT", "bo_mon": "KTPM", "role": "lecturer", "active": True},
    {"id": "u-gv2", "ma_gv": "GV002", "ho_ten": "Trần Thị Bình", "email": "gv002@dainam.edu.vn",
     "khoa": "CNTT", "bo_mon": "HTTT", "role": "lecturer", "active": True},
    {"id": "u-hd", "ma_gv": "HD001", "ho_ten": "Phạm Hội Đồng", "email": "hd@dainam.edu.vn",
     "khoa": "", "bo_mon": "", "role": "council", "active": True},
    {"id": "u-qt", "ma_gv": "QT001", "ho_ten": "Quản Trị", "email": "admin@dainam.edu.vn",
     "khoa": "", "bo_mon": "", "role": "admin", "active": True},
]


from app.security import hash_password  # noqa: E402

TEST_PW = "test123"
_PW_HASH = hash_password(TEST_PW)


@pytest.fixture()
def client():
    store = _app.state.store
    store.wipe()
    from app.rubric import get_rubric
    from app.services.ops import get_timeline

    get_rubric(store)
    get_timeline(store)
    for u in USERS:
        store.put("users", u["id"], {**u, "password_hash": _PW_HASH})
    with TestClient(_app) as c:
        yield c


@pytest.fixture()
def store():
    return _app.state.store


@pytest.fixture()
def storage():
    return _app.state.storage


def login(client: TestClient, email: str, password: str = TEST_PW) -> None:
    resp = client.post("/login", data={"login_id": email, "password": password}, follow_redirects=False)
    assert resp.status_code == 303, resp.text


DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def fill_part_a(client: TestClient, *, ten="Đề tài NCKH mẫu", loai="thuyet_minh",
                ho_ten="Nguyễn Văn An", ma_gv="GV001", khoa_bo_mon="Viện KT&KDQT",
                gvhd="PGS. TS. Vũ Hoàng Nam", thanh_vien="Trần Thị B - 2211110003") -> None:
    """Kê khai Phần A (thông tin công trình) cho người đang đăng nhập."""
    resp = client.post("/lecturer/part-a", data={
        "ten_cong_trinh": ten, "loai": loai, "ho_ten": ho_ten, "ma_gv": ma_gv,
        "khoa_bo_mon": khoa_bo_mon, "gvhd": gvhd, "thanh_vien": thanh_vien,
    }, follow_redirects=False)
    assert resp.status_code == 303, resp.text


def graded_parts_for(store, loai: str = "thuyet_minh") -> list[str]:
    from app.rubric import get_rubric, graded_parts
    return graded_parts(get_rubric(store, loai))


def upload_all_parts(client: TestClient, store, ma_gv="GV001", loai="thuyet_minh") -> None:
    """Nộp một tài liệu sản phẩm cho mỗi phần được chấm của loại công trình."""
    for part in graded_parts_for(store, loai):
        data = make_docx_bytes(f"Tài liệu phần {part}", ["Nội dung chi tiết của công trình NCKH."])
        resp = client.post(f"/lecturer/part/{part}/upload", data={"kind": "product"},
                           files={"files": (f"{ma_gv}_Phan{part}_TaiLieu.docx", data, DOCX_CT)},
                           follow_redirects=False)
        assert resp.status_code == 303, resp.text


def make_docx_bytes(title: str, lines: list[str]) -> bytes:
    import docx

    d = docx.Document()
    d.add_heading(title, level=1)
    for line in lines:
        d.add_paragraph(line)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
