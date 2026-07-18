"""Test xuất rubric ra Excel và Word."""
import io

from app.rubric import load_rubric_seed
from app.services.rubric_export import rubric_to_docx, rubric_to_xlsx
from tests.conftest import login


def _bc() -> dict:
    return load_rubric_seed()["types"]["bao_cao_co_ban"]


def test_rubric_to_xlsx_parseable():
    import openpyxl

    data = rubric_to_xlsx(_bc())
    wb = openpyxl.load_workbook(io.BytesIO(data))
    assert {"Huong_dan", "Rubric_chi_tiet", "Bang_diem", "Phan_loai"} <= set(wb.sheetnames)
    # Sheet phiếu chi tiết: đủ mã tiêu chí I/II + 4 cột mức
    rb = wb["Rubric_chi_tiet"]
    header = [c.value for c in rb[1]]
    assert "Xuất sắc (90–100%)" in header and "Chưa đạt (<50%)" in header
    crit_text = "\n".join(str(c.value) for row in rb.iter_rows() for c in row)
    for cid in ["I1", "I5", "I6", "II1", "II2"]:
        assert cid in crit_text
    # Bảng nhập điểm có công thức tổng + xếp loại
    bd_text = "\n".join(str(c.value) for row in wb["Bang_diem"].iter_rows() for c in row if c.value)
    assert "TỔNG ĐIỂM (thang 100)" in bd_text
    assert "XẾP LOẠI" in bd_text
    assert any(str(c.value).startswith("=") for row in wb["Bang_diem"].iter_rows() for c in row if c.value)


def test_rubric_to_docx_parseable():
    import docx

    data = rubric_to_docx(_bc())
    doc = docx.Document(io.BytesIO(data))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "PHIẾU ĐÁNH GIÁ" in full
    assert "Phần I" in full
    assert len(doc.tables) >= 4  # cấu trúc + 2 phần + xếp loại


def test_admin_can_download_rubric(client):
    login(client, "admin@dainam.edu.vn")
    r1 = client.get("/admin/rubric.xlsx")
    assert r1.status_code == 200
    assert "spreadsheetml" in r1.headers["content-type"]
    assert len(r1.content) > 1000
    r2 = client.get("/admin/rubric.docx")
    assert r2.status_code == 200
    assert "wordprocessingml" in r2.headers["content-type"]
    assert len(r2.content) > 1000


def test_lecturer_cannot_download_rubric(client):
    login(client, "gv001@dainam.edu.vn")
    assert client.get("/admin/rubric.xlsx", follow_redirects=False).status_code == 403


def test_admin_reload_rubric_updates_and_keeps_data(client, store):
    from app.rubric import get_rubric

    login(client, "admin@dainam.edu.vn")
    # Giả lập rubric cũ trong DB + dữ liệu khác để chứng minh không bị xóa
    store.put("config", "rubric", {"id": "rubric", "version": "0000.0", "parts": {}, "classification": []})
    store.add("submissions", {"user_id": "u-gv1", "status": "graded"})
    users_before = len(store.all("users"))
    subs_before = len(store.all("submissions"))

    r = client.post("/admin/rubric/reload", follow_redirects=False)
    assert r.status_code == 303
    # rubric đã cập nhật lên bản seed mới nhất
    assert get_rubric(store)["version"] == load_rubric_seed()["version"]
    # các dữ liệu khác còn nguyên
    assert len(store.all("users")) == users_before
    assert len(store.all("submissions")) == subs_before


def test_lecturer_cannot_reload_rubric(client):
    login(client, "gv001@dainam.edu.vn")
    assert client.post("/admin/rubric/reload", follow_redirects=False).status_code == 403


def test_download_each_of_three_rubrics(client):
    login(client, "admin@dainam.edu.vn")
    for t in ("thuyet_minh", "bao_cao_co_ban", "bao_cao_ung_dung"):
        rx = client.get(f"/admin/rubric.xlsx?type={t}")
        assert rx.status_code == 200 and "spreadsheetml" in rx.headers["content-type"] and len(rx.content) > 1000
        rd = client.get(f"/admin/rubric.docx?type={t}")
        assert rd.status_code == 200 and "wordprocessingml" in rd.headers["content-type"] and len(rd.content) > 1000


def test_rubric_json_download_has_three_types(client):
    import json

    login(client, "admin@dainam.edu.vn")
    r = client.get("/admin/rubric.json")
    assert r.status_code == 200 and "json" in r.headers["content-type"]
    data = json.loads(r.content)
    assert set(data["types"]) == {"thuyet_minh", "bao_cao_co_ban", "bao_cao_ung_dung"}


def test_rubric_upload_valid_applies(client, store):
    import json

    from app.rubric import get_rubric

    login(client, "admin@dainam.edu.vn")
    doc = json.loads(client.get("/admin/rubric.json").content)
    doc["version"] = "TEST-UPLOAD-9.9"
    payload = json.dumps(doc, ensure_ascii=False).encode("utf-8")
    r = client.post("/admin/rubric/upload",
                    files={"file": ("rubric.json", payload, "application/json")}, follow_redirects=False)
    assert r.status_code == 303
    assert get_rubric(store)["version"] == "TEST-UPLOAD-9.9"


def test_rubric_upload_invalid_rejected(client, store):
    from app.rubric import get_rubric

    login(client, "admin@dainam.edu.vn")
    before = get_rubric(store)["version"]
    # thiếu khóa 'types'
    r = client.post("/admin/rubric/upload",
                    files={"file": ("bad.json", b'{"version":"x"}', "application/json")}, follow_redirects=False)
    assert r.status_code == 303 and "rubric_msg" in r.headers["location"]
    assert get_rubric(store)["version"] == before  # không đổi
    # JSON hỏng cú pháp
    r2 = client.post("/admin/rubric/upload",
                     files={"file": ("bad2.json", b'{not json', "application/json")}, follow_redirects=False)
    assert r2.status_code == 303
    assert get_rubric(store)["version"] == before


def test_rubric_upload_download_lecturer_forbidden(client):
    login(client, "gv001@dainam.edu.vn")
    assert client.get("/admin/rubric.json", follow_redirects=False).status_code == 403
    assert client.post("/admin/rubric/upload", follow_redirects=False).status_code == 403
