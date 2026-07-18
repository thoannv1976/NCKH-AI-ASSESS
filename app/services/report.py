"""Sinh PHIẾU ĐÁNH GIÁ năng lực ứng dụng AI của giảng viên sau khi AI chấm.

Hai định dạng từ cùng một bộ dữ liệu:
- docx: python-docx (đã có sẵn).
- pdf: fpdf2 + phông DejaVuSans kèm theo (thuần Python, chạy được trên Cloud Run
  không cần thư viện hệ thống).
"""
from __future__ import annotations

import io

from app.config import GRADED_PARTS, get_settings, now_vn
from app.rubric import get_rubric
from app.services.classify import classify

_FONT_DIR = None


def _font(name: str) -> str:
    global _FONT_DIR
    if _FONT_DIR is None:
        _FONT_DIR = get_settings().base_dir / "app" / "assets" / "fonts"
    return str(_FONT_DIR / name)


def _level(score: float | None, mx: float) -> str:
    if score is None or not mx:
        return "–"
    r = score / mx
    if r >= 0.90:
        return "Xuất sắc"
    if r >= 0.70:
        return "Đạt yêu cầu"
    if r >= 0.50:
        return "Cơ bản"
    return "Chưa đạt"


def stream_reports_zip(store, subs: list[dict], fmt: str = "pdf"):
    """Đóng gói phiếu đánh giá của nhiều giảng viên thành ZIP (theo luồng).

    Bỏ qua hồ sơ chưa chấm. fmt = 'pdf' hoặc 'docx'.
    """
    import zipfile

    from app.services.downloads import _ChunkBuffer, _unique

    ext = "docx" if fmt == "docx" else "pdf"
    render = report_to_docx if fmt == "docx" else report_to_pdf
    buf = _ChunkBuffer()
    seen: set[str] = set()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED, allowZip64=True) as zf:
        for s in subs:
            ctx = build_report_context(store, s["id"])
            if not ctx:
                continue
            try:
                data = render(ctx)
            except Exception as exc:  # noqa: BLE001 — một phiếu lỗi không làm hỏng cả gói
                zf.writestr(f"{ctx['filename_base']}.LOI.txt", f"Không tạo được phiếu: {exc}")
            else:
                zf.writestr(_unique(seen, f"{ctx['filename_base']}.{ext}"), data)
            if (out := buf.take()):
                yield out
    if (out := buf.take()):
        yield out


def build_report_context(store, sid: str) -> dict | None:
    """Thu thập dữ liệu phiếu đánh giá từ điểm AI/Hội đồng. None nếu hồ sơ chưa chấm."""
    sub = store.get("submissions", sid)
    if not sub:
        return None
    owner = store.get("users", sub["user_id"]) or {}
    rubric = get_rubric(store)
    review = store.get("reviews", sid) or {}
    scores = store.find("scores", submission_id=sid)
    if not scores:
        return None
    by_id = {s["criterion_id"]: s for s in scores}

    parts, total, improvements = [], 0.0, []
    for p in GRADED_PARTS:
        pdef = rubric["parts"][p]
        crits, psum = [], 0.0
        for c in pdef["criteria"]:
            sc = by_id.get(c["id"])
            score = sc.get("final_score") if sc else None
            comment = (sc.get("council_comment") or sc.get("ai_comment") or "") if sc else ""
            crits.append({"id": c["id"], "name": c["name"], "max": c["max"],
                          "score": score, "level": _level(score, c["max"]), "comment": comment})
            if score is not None:
                psum += score
                if c["max"] and score / c["max"] < 0.5:
                    improvements.append(f"Phần {p} · {c['id']}: {comment or c['name']}")
        ptotal = round(min(psum, pdef["max_score"]), 2)
        total += ptotal
        parts.append({"key": p, "name": pdef["name"], "weight": pdef["weight"],
                      "max_score": pdef["max_score"], "total": ptotal, "criteria": crits})
    total = round(total, 2)

    final = review.get("status") in ("approved", "published")
    if final and review.get("classification"):
        cls_label, cls_key = review.get("classification_label", ""), review["classification"]
    else:
        cls = classify(total, rubric)
        cls_label, cls_key = cls["label"], cls["key"]
    cls_note = next((c["note"] for c in rubric["classification"] if c["key"] == cls_key), "")

    flags = list(dict.fromkeys(sub.get("anomaly_flags") or []))
    ranked = sorted([p for p in parts if p["max_score"]], key=lambda p: -(p["total"] / p["max_score"]))
    general = f"Hồ sơ đạt tổng {total:g}/100 điểm, xếp loại {cls_label}."
    if ranked:
        general += f" Phần đạt tỷ lệ cao nhất: {ranked[0]['name']}; phần cần chú ý: {ranked[-1]['name']}."
    if flags:
        general += f" AI phát hiện {len(flags)} dấu hiệu bất thường — đề nghị Hội đồng thẩm định kỹ."
    if not improvements:
        improvements = ["Không có tiêu chí nào dưới 50% điểm tối đa."]

    pa = sub.get("part_a") or {}
    s = get_settings()
    return {
        "ho_ten": owner.get("ho_ten") or pa.get("ho_ten", ""),
        "ma_gv": owner.get("ma_gv") or pa.get("ma_gv", ""),
        "don_vi": owner.get("khoa") or pa.get("khoa_bo_mon", ""),
        "chuc_vu": owner.get("chuc_vu", ""),
        "hoc_phan": pa.get("hoc_phan", ""),
        "cong_cu_ai": ", ".join(pa.get("cong_cu_ai") or []),
        "parts": parts, "total": total, "cls_label": cls_label, "cls_note": cls_note,
        "general": general, "improvements": improvements, "flags": flags,
        "final": final, "status": sub.get("status", ""),
        "org_name": s.org_name, "app_title": s.app_title, "program_year": s.program_year,
        "filename_base": f"PhieuDanhGia_{owner.get('ma_gv') or 'GV'}_{(owner.get('ho_ten') or '').replace(' ', '')}",
    }


# ----------------------------- DOCX -----------------------------

def report_to_docx(ctx: dict) -> bytes:
    import docx
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    def _p(text="", *, align=None, bold=False, italic=False, size=11, color=None, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold, r.italic, r.font.size = bold, italic, Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    # Quốc hiệu
    head = doc.add_table(rows=1, cols=2)
    lc, rc = head.rows[0].cells
    for cell, lines, bolds in (
        (lc, ["BỘ GIÁO DỤC VÀ ĐÀO TẠO", ctx["org_name"].upper()], [False, True]),
        (rc, ["CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", "Độc lập - Tự do - Hạnh phúc"], [True, True]),
    ):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(lines[0]).bold = bolds[0]
        q = cell.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run(lines[1]).bold = bolds[1]

    _p()
    _p("PHIẾU ĐÁNH GIÁ NĂNG LỰC ỨNG DỤNG AI CỦA GIẢNG VIÊN",
       align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15, space_after=2)
    _p(f"Chương trình {ctx['app_title']} {ctx['program_year']} · Hội đồng đánh giá cấp Trường (chấm vòng 2)",
       align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, space_after=8)

    # Thông tin
    info = doc.add_table(rows=0, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Họ và tên giảng viên", ctx["ho_ten"]),
        ("Mã giảng viên", ctx["ma_gv"]),
        ("Chức vụ", ctx["chuc_vu"]),
        ("Đơn vị / Khoa", ctx["don_vi"]),
        ("Học phần", ctx["hoc_phan"]),
        ("Công cụ AI sử dụng", ctx["cong_cu_ai"]),
        ("Tổng điểm", f"{ctx['total']:g} / 100 — {'Chính thức (Hội đồng duyệt)' if ctx['final'] else 'AI đề xuất (tạm tính)'}"),
        ("Xếp loại năng lực", ctx["cls_label"]),
    ]
    for k, v in rows:
        c0, c1 = info.add_row().cells
        c0.paragraphs[0].add_run(k).bold = True
        c1.paragraphs[0].add_run(str(v))

    _p()
    _p("BẢNG ĐIỂM CHI TIẾT", bold=True, size=12, space_after=2)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Mã", "Tiêu chí", "Tối đa", "Điểm", "Mức", "Nhận xét"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _shade(cell, hexc):
        from docx.oxml.ns import qn
        sh = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {qn("w:fill"): hexc})
        cell._tc.get_or_add_tcPr().append(sh)

    for p in ctx["parts"]:
        r = tbl.add_row().cells
        r[0].merge(r[5])
        run = r[0].paragraphs[0].add_run(
            f"PHẦN {p['key']} – {p['name']} (trọng số {p['weight']}%, tối đa {p['max_score']} điểm)")
        run.bold = True
        _shade(r[0], "FFEDD5")
        for c in p["criteria"]:
            cells = tbl.add_row().cells
            vals = [c["id"], c["name"], f"{c['max']:g}",
                    (f"{c['score']:g}" if c["score"] is not None else "–"), c["level"], c["comment"]]
            for i, v in enumerate(vals):
                cells[i].paragraphs[0].add_run(v)
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9)
        rr = tbl.add_row().cells
        rr[0].merge(rr[2])
        rr[0].paragraphs[0].add_run(f"Cộng Phần {p['key']}").bold = True
        rr[3].paragraphs[0].add_run(f"{p['total']:g}").bold = True
        rr[4].merge(rr[5])

    tr = tbl.add_row().cells
    tr[0].merge(tr[2])
    tr[0].paragraphs[0].add_run("TỔNG ĐIỂM (thang 100)").bold = True
    tr[3].paragraphs[0].add_run(f"{ctx['total']:g}").bold = True
    tr[4].merge(tr[5])
    tr[4].paragraphs[0].add_run(ctx["cls_label"]).bold = True

    _p()
    _p("NHẬN XÉT CHUNG", bold=True, size=12, space_after=2)
    _p(ctx["general"])
    _p("ĐIỂM CẦN HOÀN THIỆN", bold=True, size=12, space_after=2)
    for it in ctx["improvements"]:
        doc.add_paragraph(it, style="List Bullet")
    if ctx["flags"]:
        _p("DẤU HIỆU BẤT THƯỜNG DO AI PHÁT HIỆN", bold=True, size=12, color=RGBColor(0xC0, 0x39, 0x2B), space_after=2)
        for f in ctx["flags"]:
            doc.add_paragraph(f, style="List Bullet")
    _p("ĐỊNH HƯỚNG SỬ DỤNG KẾT QUẢ", bold=True, size=12, space_after=2)
    _p(f"{ctx['cls_label']}: {ctx['cls_note']}")
    _p("Kết quả dùng để phát triển đội ngũ, KHÔNG dùng xử lý thi đua/kỷ luật. "
       "Giảng viên được phản hồi trong 03 ngày làm việc kể từ ngày công bố.",
       italic=True, size=10, color=RGBColor(0x6B, 0x72, 0x80))

    _p()
    d = now_vn()
    _p(f"Ngày {d:%d} tháng {d:%m} năm {d:%Y}", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, size=10)
    sig = doc.add_table(rows=1, cols=2)
    for cell, title in zip(sig.rows[0].cells,
                           ["THÀNH VIÊN HỘI ĐỒNG CHẤM", "CHỦ TỊCH HỘI ĐỒNG ĐÁNH GIÁ"]):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(title).bold = True
        sub = cell.add_paragraph("(Ký, ghi rõ họ tên)")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------- PDF (fpdf2) -----------------------------

def report_to_pdf(ctx: dict) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    pdf.add_font("dj", "", _font("DejaVuSans.ttf"))
    pdf.add_font("dj", "B", _font("DejaVuSans-Bold.ttf"))

    def text(t, size=11, b=False, align="L", h=6, fill=False):
        pdf.set_font("dj", "B" if b else "", size)
        pdf.multi_cell(0, h, t, align=align, new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=fill)

    # Quốc hiệu (2 cột)
    pdf.set_font("dj", "", 9)
    y0 = pdf.get_y()
    w = (pdf.w - pdf.l_margin - pdf.r_margin) / 2
    pdf.multi_cell(w, 5, f"BỘ GIÁO DỤC VÀ ĐÀO TẠO\n{ctx['org_name'].upper()}", align="C",
                   new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.set_xy(pdf.l_margin + w, y0)
    pdf.multi_cell(w, 5, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc",
                   align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)
    text("PHIẾU ĐÁNH GIÁ NĂNG LỰC ỨNG DỤNG AI CỦA GIẢNG VIÊN", size=14, b=True, align="C", h=7)
    text(f"Chương trình {ctx['app_title']} {ctx['program_year']} · Hội đồng đánh giá cấp Trường (chấm vòng 2)",
         size=9, align="C", h=5)
    pdf.ln(2)

    # Thông tin
    info = [("Họ và tên giảng viên", ctx["ho_ten"]), ("Mã giảng viên", ctx["ma_gv"]),
            ("Chức vụ", ctx["chuc_vu"]), ("Đơn vị / Khoa", ctx["don_vi"]),
            ("Học phần", ctx["hoc_phan"]), ("Công cụ AI sử dụng", ctx["cong_cu_ai"]),
            ("Tổng điểm", f"{ctx['total']:g} / 100 ({'chính thức' if ctx['final'] else 'AI đề xuất'})"),
            ("Xếp loại năng lực", ctx["cls_label"])]
    lw = 46
    for k, v in info:
        pdf.set_font("dj", "B", 10)
        yk = pdf.get_y()
        pdf.multi_cell(lw, 6, k, border=1, align="L", new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_xy(pdf.l_margin + lw, yk)
        pdf.set_font("dj", "", 10)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - lw, 6, str(v) or "—", border=1,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)
    text("BẢNG ĐIỂM CHI TIẾT", size=11, b=True, h=6)

    full = pdf.w - pdf.l_margin - pdf.r_margin

    def crit_table(part):
        widths = (14, 52, 14, 14, 22, full - 116)
        pdf.set_fill_color(255, 237, 213)
        text(f"PHẦN {part['key']} – {part['name']} (trọng số {part['weight']}%, tối đa {part['max_score']} điểm)",
             size=9, b=True, h=6, fill=True)
        with pdf.table(col_widths=widths, line_height=5, first_row_as_headings=True,
                       text_align=("CENTER", "LEFT", "CENTER", "CENTER", "CENTER", "LEFT")) as table:
            table.row(("Mã", "Tiêu chí", "Tối đa", "Điểm", "Mức", "Nhận xét"))
            for c in part["criteria"]:
                table.row((c["id"], c["name"], f"{c['max']:g}",
                           (f"{c['score']:g}" if c["score"] is not None else "–"),
                           c["level"], c["comment"] or ""))
        pdf.set_font("dj", "B", 9)
        pdf.multi_cell(0, 5, f"Cộng Phần {part['key']}: {part['total']:g}/{part['max_score']:g}",
                       align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    for p in ctx["parts"]:
        crit_table(p)

    pdf.ln(1)
    text(f"TỔNG ĐIỂM (thang 100): {ctx['total']:g}  —  Xếp loại: {ctx['cls_label']}", size=11, b=True, h=7)
    pdf.ln(1)
    text("NHẬN XÉT CHUNG", size=11, b=True, h=6)
    text(ctx["general"], size=10, h=5)
    text("ĐIỂM CẦN HOÀN THIỆN", size=11, b=True, h=6)
    for it in ctx["improvements"]:
        text(f"•  {it}", size=10, h=5)
    if ctx["flags"]:
        pdf.set_text_color(192, 57, 43)
        text("DẤU HIỆU BẤT THƯỜNG DO AI PHÁT HIỆN", size=11, b=True, h=6)
        for f in ctx["flags"]:
            text(f"•  {f}", size=10, h=5)
        pdf.set_text_color(0, 0, 0)
    text("ĐỊNH HƯỚNG SỬ DỤNG KẾT QUẢ", size=11, b=True, h=6)
    text(f"{ctx['cls_label']}: {ctx['cls_note']}", size=10, h=5)
    pdf.set_text_color(110, 114, 128)
    text("Kết quả dùng để phát triển đội ngũ, KHÔNG dùng xử lý thi đua/kỷ luật. "
         "Giảng viên được phản hồi trong 03 ngày làm việc kể từ ngày công bố.", size=9, h=5)
    pdf.set_text_color(0, 0, 0)

    pdf.ln(4)
    d = now_vn()
    text(f"Ngày {d:%d} tháng {d:%m} năm {d:%Y}", size=10, align="R", h=5)
    y = pdf.get_y()
    half = full / 2
    pdf.set_font("dj", "B", 10)
    pdf.multi_cell(half, 6, "THÀNH VIÊN HỘI ĐỒNG CHẤM", align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.set_xy(pdf.l_margin + half, y)
    pdf.multi_cell(half, 6, "CHỦ TỊCH HỘI ĐỒNG ĐÁNH GIÁ", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("dj", "", 9)
    y = pdf.get_y()
    pdf.multi_cell(half, 5, "(Ký, ghi rõ họ tên)", align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.set_xy(pdf.l_margin + half, y)
    pdf.multi_cell(half, 5, "(Ký, ghi rõ họ tên)", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    out = pdf.output()
    return bytes(out)
