"""Xuất rubric chấm điểm ra Excel (.xlsx) và Word (.docx).

Bản Excel là BỘ RUBRIC CHUẨN CHO HỘI ĐỒNG (chấm vòng 2): gồm hướng dẫn,
rubric 4 mức chi tiết, bảng nhập điểm tự tính tổng + xếp loại, và bảng phân loại.
"""
from __future__ import annotations

import io

from app.config import get_settings

LEVELS = [
    ("xuat_sac", "Xuất sắc (90–100%)"),
    ("dat", "Đạt yêu cầu (70–89%)"),
    ("co_ban", "Cơ bản (50–69%)"),
    ("chua_dat", "Chưa đạt (<50%)"),
]


def _graded_parts(rubric: dict):
    return [(k, p) for k, p in rubric["parts"].items() if p.get("graded")]


def _level_range(rubric: dict, i: int) -> str:
    levels = sorted(rubric["classification"], key=lambda x: -x["min"])
    if i == 0:
        return f"{levels[0]['min']}–100"
    if i == len(levels) - 1:
        return f"Dưới {levels[i - 1]['min']}"
    return f"{levels[i]['min']}–{levels[i - 1]['min'] - 1}"


def rubric_to_xlsx(rubric: dict) -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    head_fill = PatternFill("solid", fgColor="EA580C")
    sub_fill = PatternFill("solid", fgColor="FFEDD5")
    input_fill = PatternFill("solid", fgColor="FEF9C3")  # ô vàng Hội đồng nhập
    head_font = Font(bold=True, color="FFFFFF")
    bold = Font(bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center")
    thin = Border(*[Side(style="thin", color="D1D5DB")] * 4)

    s = get_settings()
    wb = openpyxl.Workbook()

    # ---------- Sheet 1: Hướng dẫn ----------
    ws = wb.active
    ws.title = "Huong_dan"
    lines = [
        (f"RUBRIC CHẤM ĐIỂM – ĐÁNH GIÁ NĂNG LỰC ỨNG DỤNG AI CỦA GIẢNG VIÊN {s.org_short} {s.program_year}", 14, True),
        (f"Phiên bản rubric: {rubric.get('version', '')} · Dùng cho Hội đồng đánh giá cấp Trường (chấm vòng 2)", 11, False),
        ("", 11, False),
        ("1. Thang điểm: tổng 100 điểm theo trọng số các Phần B–G. Phần A là điều kiện hợp lệ (Đạt/Không đạt), không tính điểm.", 11, False),
        ("2. Mỗi tiêu chí chấm theo 4 mức: Xuất sắc (90–100% điểm), Đạt yêu cầu (70–89%), Cơ bản (50–69%), Chưa đạt (<50%).", 11, False),
        ("   Hội đồng xác định mức phù hợp nhất với bài làm rồi cho điểm trong khoảng của mức đó (xem sheet 'Rubric_chi_tiet').", 11, False),
        (f"3. Phần mềm {s.app_title} đã chấm vòng 1: mỗi tiêu chí 2 lượt độc lập, chênh >15% chấm lượt 3 lấy trung vị.", 11, False),
        ("   Cột 'Điểm AI đề xuất' ở sheet 'Bang_diem' là kết quả vòng 1 để Hội đồng tham khảo.", 11, False),
        ("4. Quy tắc minh chứng: sản phẩm thiếu minh chứng AI bị trừ tối đa 50% điểm của tiêu chí liên quan.", 11, False),
        ("5. Phần E có điểm thưởng E_BONUS (tối đa +2) cho nhiệm vụ khuyến khích, nhưng tổng Phần E không vượt 20 điểm.", 11, False),
        ("6. Hội đồng nhập điểm và nhận xét vào CỘT VÀNG ở sheet 'Bang_diem'; tổng điểm và xếp loại tự động tính.", 11, False),
        ("7. Điểm cuối cùng là điểm Hội đồng phê duyệt. Kết quả dùng để phát triển đội ngũ, KHÔNG dùng xử lý thi đua/kỷ luật.", 11, False),
        ("", 11, False),
        ("Các sheet: Rubric_chi_tiet (rubric 4 mức) · Bang_diem (nhập điểm, tự tính) · Phan_loai (ngưỡng xếp loại).", 11, False),
    ]
    for i, (text, size, b) in enumerate(lines, 1):
        ws.cell(row=i, column=1, value=text).font = Font(bold=b or size > 12, size=size)
    ws.column_dimensions["A"].width = 120

    # ---------- Sheet 2: Rubric chi tiết (4 mức) ----------
    ws2 = wb.create_sheet("Rubric_chi_tiet")
    headers = ["Phần", "Trọng số", "Mã", "Tiêu chí", "Điểm tối đa",
               "Xuất sắc (90–100%)", "Đạt yêu cầu (70–89%)", "Cơ bản (50–69%)", "Chưa đạt (<50%)"]
    ws2.append(headers)
    for c in ws2[1]:
        c.fill, c.font, c.alignment, c.border = head_fill, head_font, center, thin
    for k, p in _graded_parts(rubric):
        first = True
        for crit in p["criteria"]:
            lv = crit.get("levels", {})
            row = [k if first else "", f"{p['weight']}%" if first else "", crit["id"],
                   crit["name"], crit["max"],
                   lv.get("xuat_sac", ""), lv.get("dat", ""), lv.get("co_ban", ""), lv.get("chua_dat", "")]
            ws2.append(row)
            for ci, cell in enumerate(ws2[ws2.max_row], 1):
                cell.alignment = wrap if ci in (4, 6, 7, 8, 9) else center
                cell.border = thin
            first = False
    widths = [6, 8, 10, 38, 8, 30, 28, 26, 24]
    for i, w in enumerate(widths, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"

    # ---------- Sheet 3: Bảng nhập điểm (tự tính tổng + xếp loại) ----------
    ws3 = wb.create_sheet("Bang_diem")
    info = [(f"BẢNG NHẬP ĐIỂM HỘI ĐỒNG – {s.app_title.upper()} {s.program_year}", ""),
            ("Họ tên GV:", ""), ("Mã GV:", ""), ("Đơn vị/Bộ môn:", ""),
            ("Học phần:", ""), ("Thành viên Hội đồng chấm:", "")]
    for label, val in info:
        ws3.append([label, val])
    ws3["A1"].font = Font(bold=True, size=13)
    ws3.append([])
    hdr_row = ws3.max_row + 1
    ws3.append(["Mã", "Tiêu chí", "Điểm tối đa", "Điểm AI đề xuất", "Điểm Hội đồng", "Nhận xét"])
    for c in ws3[hdr_row]:
        c.fill, c.font, c.alignment, c.border = head_fill, head_font, center, thin

    part_subtotal_rows: list[tuple[str, int, int]] = []  # (part, first_data_row, last_data_row)
    for k, p in _graded_parts(rubric):
        ws3.append([f"PHẦN {k}", f"{p['name']}  (trọng số {p['weight']}%)", "", "", "", ""])
        prow = ws3[ws3.max_row]
        for c in prow:
            c.fill, c.font = sub_fill, bold
        first_data = ws3.max_row + 1
        for crit in p["criteria"]:
            ws3.append([crit["id"], crit["name"], crit["max"], "", "", ""])
            r = ws3[ws3.max_row]
            r[1].alignment = wrap
            r[4].fill = input_fill  # Điểm Hội đồng — ô vàng nhập
            r[5].fill = input_fill  # Nhận xét — ô vàng nhập
            for cell in r:
                cell.border = thin
        last_data = ws3.max_row
        # Dòng cộng phần (Phần E kẹp trần max_score)
        cmax = get_column_letter(5)
        sum_expr = f"SUM({cmax}{first_data}:{cmax}{last_data})"
        formula = f"=MIN({sum_expr},{p['max_score']})" if k == "E" else f"={sum_expr}"
        ws3.append([f"Cộng Phần {k}", "", p["max_score"], "", None, ""])
        srow = ws3[ws3.max_row]
        srow[0].font = bold
        ws3.cell(row=ws3.max_row, column=5, value=formula).font = bold
        part_subtotal_rows.append((k, first_data, last_data))
        sub_row_idx = ws3.max_row
        part_subtotal_rows[-1] = (k, sub_row_idx, sub_row_idx)  # lưu dòng cộng phần

    # Tổng điểm = tổng các dòng "Cộng Phần"
    total_row = ws3.max_row + 2
    sub_cells = "+".join(f"E{r}" for _, r, _ in part_subtotal_rows)
    ws3.cell(row=total_row, column=2, value="TỔNG ĐIỂM (thang 100)").font = Font(bold=True, size=12)
    ws3.cell(row=total_row, column=3, value=100)
    ws3.cell(row=total_row, column=5, value=f"={sub_cells}").font = Font(bold=True, size=12)
    # Xếp loại tự động theo tổng (cột E dòng total)
    levels = sorted(rubric["classification"], key=lambda x: -x["min"])
    tot = f"E{total_row}"
    cls = (f'=IF({tot}>={levels[0]["min"]},"{levels[0]["label"]}",'
           f'IF({tot}>={levels[1]["min"]},"{levels[1]["label"]}",'
           f'IF({tot}>={levels[2]["min"]},"{levels[2]["label"]}","{levels[3]["label"]}")))')
    ws3.cell(row=total_row + 1, column=2, value="XẾP LOẠI NĂNG LỰC").font = Font(bold=True, size=12)
    ws3.cell(row=total_row + 1, column=5, value=cls).font = Font(bold=True, size=12, color="EA580C")
    for w, col in zip([10, 46, 11, 14, 14, 50], "ABCDEF"):
        ws3.column_dimensions[col].width = w
    ws3.freeze_panes = "A" + str(hdr_row + 1)

    # ---------- Sheet 4: Phân loại ----------
    ws4 = wb.create_sheet("Phan_loai")
    ws4.append(["XẾP LOẠI NĂNG LỰC VÀ ĐỊNH HƯỚNG SỬ DỤNG KẾT QUẢ"])
    ws4["A1"].font = Font(bold=True, size=13)
    ws4.append([])
    ws4.append(["Mức năng lực", "Khoảng điểm", "Định hướng sử dụng kết quả"])
    for c in ws4[3]:
        c.fill, c.font, c.border = head_fill, head_font, thin
    for i, lv in enumerate(levels):
        ws4.append([lv["label"], _level_range(rubric, i), lv.get("note", "")])
        ws4[ws4.max_row][2].alignment = wrap
    ws4.append([])
    ws4.append(["Lưu ý: Kết quả dùng để phát triển đội ngũ, KHÔNG dùng xử lý thi đua/kỷ luật. "
                "Giảng viên được phản hồi trong 03 ngày làm việc kể từ ngày công bố."])
    for w, col in zip([24, 14, 95], "ABC"):
        ws4.column_dimensions[col].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def rubric_to_docx(rubric: dict) -> bytes:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    s = get_settings()
    doc = docx.Document()
    doc.add_heading("RUBRIC CHẤM ĐIỂM – ĐÁNH GIÁ NĂNG LỰC ỨNG DỤNG AI", level=0)
    p = doc.add_paragraph(f"{s.org_name} ({s.org_short}) — Năm {s.program_year} · "
                          "Dùng cho Hội đồng đánh giá cấp Trường (chấm vòng 2)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Phiên bản rubric: {rubric.get('version', '')} · Thang điểm 100").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Cấu trúc và trọng số", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Phần", "Nội dung", "Trọng số", "Điểm"]):
        t.rows[0].cells[i].text = h
    for k, pp in rubric["parts"].items():
        graded = pp.get("graded")
        row = t.add_row().cells
        row[0].text, row[1].text = k, pp["name"]
        row[2].text = f"{pp['weight']}%" if graded else "–"
        row[3].text = str(pp["max_score"]) if graded else "Đạt/Không đạt"

    for k, pp in _graded_parts(rubric):
        doc.add_heading(f"Phần {k} — {pp['name']} (trọng số {pp['weight']}%, tối đa {pp['max_score']} điểm)", level=1)
        doc.add_paragraph(pp.get("description", ""))
        if pp.get("products"):
            doc.add_paragraph("Sản phẩm phải nộp:").runs[0].bold = True
            for pr in pp["products"]:
                doc.add_paragraph(pr, style="List Bullet")
        ct = doc.add_table(rows=1, cols=6)
        ct.style = "Light Grid Accent 1"
        for i, h in enumerate(["Mã", "Tiêu chí", "Điểm", "Xuất sắc", "Đạt", "Cơ bản / Chưa đạt"]):
            ct.rows[0].cells[i].text = h
        for crit in pp["criteria"]:
            lv = crit.get("levels", {})
            row = ct.add_row().cells
            row[0].text = crit["id"]
            row[1].text = crit["name"]
            row[2].text = str(crit["max"])
            row[3].text = lv.get("xuat_sac", "")
            row[4].text = lv.get("dat", "")
            row[5].text = f"Cơ bản: {lv.get('co_ban', '')}\nChưa đạt: {lv.get('chua_dat', '')}"
        for r in ct.rows:
            for cell in r.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_heading("Xếp loại năng lực", level=1)
    lt = doc.add_table(rows=1, cols=3)
    lt.style = "Light Grid Accent 1"
    for i, h in enumerate(["Mức năng lực", "Điểm", "Định hướng sử dụng kết quả"]):
        lt.rows[0].cells[i].text = h
    levels = sorted(rubric["classification"], key=lambda x: -x["min"])
    for i, lv in enumerate(levels):
        row = lt.add_row().cells
        row[0].text, row[1].text, row[2].text = lv["label"], _level_range(rubric, i), lv.get("note", "")

    note = doc.add_paragraph("Lưu ý: Kết quả dùng để phát triển đội ngũ, KHÔNG dùng xử lý thi đua/kỷ luật.")
    note.runs[0].font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
