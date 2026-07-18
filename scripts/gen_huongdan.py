"""Sinh HƯỚNG DẪN SỬ DỤNG (.docx) cho DNU AI-Assess, kèm ảnh minh họa giao diện.

Ảnh được render từ HTML/CSS tái hiện đúng màn hình thật (header, menu, thẻ, bảng, nút —
đúng nhãn, bố cục, màu sắc) qua WeasyPrint → PDF → PNG (pymupdf), tự cắt nền thừa (Pillow).
Đây là ẢNH MINH HỌA GIAO DIỆN; có thể thay bằng ảnh chụp thực tế trên hệ thống đã triển khai.

Chạy: python scripts/gen_huongdan.py [out.docx]
"""
from __future__ import annotations

import sys
import tempfile
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

import fitz  # pymupdf
import weasyprint
from PIL import Image

IMGDIR = Path(tempfile.mkdtemp(prefix="dnu-hd-"))

THEME = """
@page { size: 1180px 1500px; margin: 0; }
* { box-sizing: border-box; }
body { margin: 0; background: #eef2f6; font-family: 'DejaVu Sans', sans-serif; color: #1e293b; font-size: 15px; }
.topbar { background: #c2410c; color: #fff; padding: 14px 22px; }
.topbar .brand { font-size: 19px; font-weight: bold; }
.topbar .sub { font-size: 12px; color: #fed7aa; font-weight: normal; margin-left: 10px; }
.topbar .nav { text-align: right; font-size: 13px; color: #ffe7d6; }
.topbar .nav span { margin-left: 14px; }
.topbar .nav .me { background: #9a3412; padding: 4px 8px; border-radius: 5px; color: #fff; }
.wrap { padding: 22px 26px; }
.h1 { font-size: 21px; font-weight: bold; margin: 0 0 14px; }
.muted { color: #64748b; }
.card { background: #fff; border-radius: 12px; box-shadow: 0 1px 3px rgba(0,0,0,.12); padding: 18px 20px; margin-bottom: 16px; }
.card h2 { font-size: 16px; margin: 0 0 10px; }
.row { display: flex; gap: 16px; }
.col { flex: 1; }
.btn { display: inline-block; color: #fff; border-radius: 8px; padding: 9px 16px; font-weight: bold; font-size: 14px; }
.btn-org { background: #ea580c; } .btn-slate { background: #334155; } .btn-green { background: #16a34a; }
.btn-blue { background: #2563eb; } .btn-sm { padding: 5px 10px; font-size: 12px; }
.pill { display: inline-block; border-radius: 999px; padding: 3px 10px; font-size: 12px; }
.pill-green { background: #d1fae5; color: #047857; } .pill-amber { background: #fef3c7; color: #b45309; }
.pill-slate { background: #e2e8f0; color: #475569; } .pill-red { background: #fee2e2; color: #dc2626; }
table { width: 100%; border-collapse: collapse; font-size: 13px; }
th { text-align: left; color: #94a3b8; font-size: 11px; border-bottom: 1px solid #e2e8f0; padding: 7px 8px; text-transform: uppercase; }
td { border-bottom: 1px solid #eef2f6; padding: 8px; }
.t-head { background: #c2410c; } .t-head th { color: #fff; }
.in { border: 1px solid #cbd5e1; border-radius: 8px; padding: 9px 11px; color: #94a3b8; font-size: 13px; background: #fff; margin-top: 4px; }
.lbl { font-size: 13px; font-weight: bold; }
.note { background: #fff7ed; border: 1px solid #fed7aa; color: #9a3412; border-radius: 8px; padding: 10px 12px; font-size: 13px; }
.big { font-size: 34px; font-weight: 800; color: #c2410c; }
.center { text-align: center; }
.stat { background:#fff; border-radius:12px; box-shadow:0 1px 3px rgba(0,0,0,.12); padding:14px; text-align:center; }
"""

NAVS = {
    "lecturer": ["Hồ sơ của tôi", "Kết quả"],
    "council": ["Thẩm định", "Phản hồi", "Bảng điểm", "Báo cáo"],
    "admin": ["Vận hành", "Người dùng", "Cấu hình AI", "Tải sản phẩm", "Thẩm định", "Bảng điểm", "Báo cáo"],
}
ME = {"lecturer": "Nguyễn Văn An (Giảng viên)", "council": "Phạm Hội Đồng (Hội đồng)",
      "admin": "Quản trị viên (Quản trị)"}


def topbar(role: str | None) -> str:
    if role is None:
        nav = ""
    else:
        items = "".join(f"<span>{x}</span>" for x in NAVS[role])
        nav = f"{items}<span class='me'>{ME[role]}</span><span>Đổi mật khẩu</span><span>Thoát</span>"
    return (
        "<div class='topbar'><table><tr>"
        "<td class='brand'>DNU AI-Assess<span class='sub'>Đánh giá năng lực ứng dụng AI của giảng viên — 2026</span></td>"
        f"<td class='nav'>{nav}</td></tr></table></div>"
    )


def screen(role: str | None, body: str) -> str:
    return f"<html><head><meta charset='utf-8'><style>{THEME}</style></head><body>{topbar(role)}<div class='wrap'>{body}</div></body></html>"


_EMOJI = "⬇⬆⏰⏳⚡⚠📧👥🔑💰📦📊💾🎓🖼"


def render(name: str, html: str) -> Path:
    html = html.translate({ord(c): None for c in _EMOJI}).replace("  ", " ")
    pdf = weasyprint.HTML(string=html).write_pdf()
    doc = fitz.open(stream=pdf, filetype="pdf")
    pix = doc[0].get_pixmap(matrix=fitz.Matrix(1.6, 1.6))
    raw = IMGDIR / (name + "_raw.png")
    pix.save(str(raw))
    # cắt nền thừa phía dưới
    im = Image.open(raw).convert("RGB")
    w, h = im.size
    px = im.load()

    def has(y: int) -> bool:
        for x in range(0, w, 6):
            r, g, b = px[x, y]
            if not (235 <= r <= 245 and 240 <= g <= 250 and 244 <= b <= 252):
                return True
        return False

    last = h - 1
    for y in range(h - 1, -1, -1):
        if has(y):
            last = y
            break
    out = IMGDIR / (name + ".png")
    im.crop((0, 0, w, min(h, last + 28))).save(out)
    return out


# ---------- thành phần dựng nhanh ----------

def tbl(headers, rows, head_class="t-head"):
    h = "".join(f"<th>{x}</th>" for x in headers)
    body = ""
    for r in rows:
        body += "<tr>" + "".join(f"<td>{c}</td>" for c in r) + "</tr>"
    return f"<table><tr class='{head_class}'>{h}</tr>{body}</table>"


def stats(items):
    cells = "".join(
        f"<td style='padding:0 8px'><div class='stat'><div class='big'>{v}</div>"
        f"<div class='muted' style='font-size:12px'>{k}</div></div></td>" for k, v in items)
    return f"<table><tr>{cells}</tr></table>"


# ---------- định nghĩa các màn hình ----------

def build_screens() -> dict[str, Path]:
    s: dict[str, Path] = {}

    # 1. Đăng nhập
    s["login"] = render("login", screen(None,
        "<div style='max-width:430px;margin:30px auto'><div class='card center'>"
        "<div style='font-size:22px;font-weight:bold;color:#c2410c'>DNU AI-Assess</div>"
        "<div class='muted' style='font-size:13px;margin:8px 0 18px'>Hệ thống đánh giá năng lực ứng dụng AI của giảng viên<br>Trường Đại học Đại Nam — 2026</div>"
        "<div style='text-align:left'><div class='lbl'>ID đăng nhập</div>"
        "<div class='in'>Email hoặc mã giảng viên (vd: gv001@dainam.edu.vn hoặc GV001)</div>"
        "<div class='lbl' style='margin-top:10px'>Mật khẩu</div><div class='in'>••••••••</div>"
        "<div class='btn btn-org' style='width:100%;text-align:center;margin-top:14px'>Đăng nhập</div></div>"
        "<div class='muted' style='font-size:12px;margin-top:12px'>Quên mật khẩu hoặc chưa có tài khoản? Liên hệ quản trị viên / Ban tổ chức.</div>"
        "</div></div>"))

    # 2. Giảng viên — trang hồ sơ
    parts = [("A", "Thông tin chung", "Điều kiện hợp lệ", "Đầy đủ", "pill-green"),
             ("B", "Phát triển CTĐT và học phần", "20 điểm (20%)", "Đủ thành phần", "pill-green"),
             ("C", "Giảng dạy và hỗ trợ học tập", "40 điểm (40%)", "Đủ thành phần", "pill-green"),
             ("D", "Đánh giá người học", "10 điểm (10%)", "Chưa đủ", "pill-slate"),
             ("E", "Nghiên cứu khoa học", "20 điểm (20%)", "Đủ thành phần", "pill-green"),
             ("F", "Phục vụ cộng đồng, doanh nghiệp", "5 điểm (5%)", "Đủ thành phần", "pill-green")]
    cells = []
    for code, name, pts, st, cls in parts:
        cells.append(
            f"<td style='width:50%;padding:6px;vertical-align:top'><div class='card' style='margin:0'>"
            f"<b>Phần {code} — {name}</b> <span class='muted' style='font-size:12px'>· {pts}</span><br>"
            f"<span style='font-size:13px'>Sản phẩm: <b>2</b> &nbsp; Minh chứng AI: <b>1</b></span> "
            f"<span class='pill {cls}'>{st}</span><br>"
            f"<span style='color:#c2410c;font-size:13px;font-weight:bold'>Nộp sản phẩm &amp; minh chứng →</span></div></td>")
    rows = "".join(f"<tr>{cells[i]}{cells[i + 1]}</tr>" for i in range(0, len(cells), 2))
    s["lec_dashboard"] = render("lec_dashboard", screen("lecturer",
        "<div class='h1'>Hồ sơ đánh giá năng lực AI — Nguyễn Văn An <span class='pill pill-green'>Đã nộp</span></div>"
        "<div class='muted' style='margin-bottom:8px'>Hạn nộp: <b>17h00 ngày 30/6/2026</b></div>"
        f"<table>{rows}</table>"
        "<div class='center' style='margin-top:14px'><span class='btn btn-org'>NỘP HỒ SƠ</span></div>"))

    # 3. Giảng viên — nộp 1 phần
    s["lec_part"] = render("lec_part", screen("lecturer",
        "<div class='muted' style='font-size:13px'>← Quay lại hồ sơ</div>"
        "<div class='h1'>Phần C — Giảng dạy và hỗ trợ học tập</div>"
        "<div class='row'><div class='col'><div class='card'><b>Yêu cầu của phần</b>"
        "<div class='muted' style='font-size:13px;margin-top:6px'>C1 Giáo án tích hợp AI · C2 Học liệu số (slide ≥15 trang) · "
        "C3 Trợ lý AI/bộ prompt · C4 Dạy học phân hóa. Tên tệp: GV001_PhanC_TênSảnPhẩm. Định dạng docx/pdf/pptx/xlsx/mp4, tối đa 200MB.</div></div></div>"
        "<div class='col'>"
        "<div class='card'><b>Sản phẩm</b>"
        "<div style='border:1px solid #e2e8f0;border-radius:8px;padding:8px;margin:8px 0;font-size:13px'>GV001_PhanC_KeHoachBaiGiang.docx &nbsp; <span class='muted'>0.1 MB</span> &nbsp; <span style='color:#dc2626'>Xóa</span></div>"
        "<div style='border:2px dashed #cbd5e1;border-radius:8px;padding:10px'><div class='muted' style='font-size:11px;font-weight:bold'>TẢI TỆP LÊN (chọn được nhiều tệp)</div>"
        "<div class='in'>Chọn tệp...</div><div class='btn btn-slate btn-sm' style='margin-top:6px'>Tải lên</div></div></div>"
        "<div class='card'><b>Minh chứng sử dụng AI</b>"
        "<div style='border:2px dashed #cbd5e1;border-radius:8px;padding:10px;margin-top:8px'><div class='muted' style='font-size:11px;font-weight:bold'>HOẶC DÁN LIÊN KẾT (chatbot, hội thoại AI, video...)</div>"
        "<div class='in'>https://...</div><div class='btn btn-slate btn-sm' style='margin-top:6px'>Thêm liên kết</div></div></div>"
        "</div></div>"))

    # 4. Giảng viên — kết quả
    s["lec_result"] = render("lec_result", screen("lecturer",
        "<div class='h1'>Kết quả đánh giá năng lực ứng dụng AI 2026</div>"
        + stats([("Tổng điểm", "78/100"), ("Mức năng lực", "Thành thạo"), ("Phản hồi", "Còn trong hạn")])
        + "<div class='card' style='margin-top:14px'><b>Phần C — Giảng dạy và hỗ trợ học tập</b> "
        "<span style='float:right;color:#c2410c;font-weight:bold'>31/40</span>"
        + tbl(["Tiêu chí", "Điểm", "Nhận xét"], [
            ["C1", "8/10", "Giáo án rõ mục tiêu, tích hợp AI 2 hoạt động hợp lý..."],
            ["C2", "8/10", "Slide đạt yêu cầu, học liệu bổ trợ khá..."],
            ["C3", "8/10", "Bộ prompt hữu ích, đã thử nghiệm..."],
            ["C4", "7/10", "Phân hóa 2 nhóm, vai trò AI nêu được nhưng ngắn..."],
        ]) + "</div>"
        "<div class='note'>Gửi phản hồi về kết quả trong 03 ngày làm việc kể từ ngày công bố.</div>"))

    # 5. Hội đồng — danh sách thẩm định
    s["council_list"] = render("council_list", screen("council",
        "<div class='h1'>Thẩm định hồ sơ</div>"
        + tbl(["Giảng viên", "Đơn vị", "Hợp lệ", "Điểm AI", "Thẩm định bắt buộc", "Trạng thái", ""], [
            ["<b>Lê Minh Châu</b> GV003", "CNTT", "✓", "<b>88</b> ✓AI", "<span class='pill pill-amber'>⚠ Bắt buộc</span>", "Đã chấm", "<b style='color:#c2410c'>Thẩm định →</b>"],
            ["<b>Nguyễn Văn An</b> GV001", "CNTT", "✓", "<b>78</b> ✓AI", "–", "Đã chấm", "<b style='color:#c2410c'>Thẩm định →</b>"],
            ["<b>Trần Thị Bình</b> GV002", "CNTT", "✓", "<b>65</b> ✓AI", "–", "Đã chấm", "<b style='color:#c2410c'>Thẩm định →</b>"],
            ["<b>Phạm Quốc Dũng</b> GV004", "QTKD", "✓", "<span class='muted'>đang chấm</span>", "–", "Đã khóa", "<b style='color:#c2410c'>Thẩm định →</b>"],
        ])))

    # 6. Hội đồng — chi tiết chấm
    s["council_detail"] = render("council_detail", screen("council",
        "<div class='muted' style='font-size:13px'>← Danh sách hồ sơ</div>"
        "<div class='h1'>Nguyễn Văn An <span class='muted' style='font-size:15px'>(GV001 — CNTT · Giảng viên)</span> "
        "<span class='pill pill-slate'>graded</span> <span style='float:right'>Tổng hiện tại: <b style='color:#c2410c'>78</b>/100</span></div>"
        "<div class='card' style='background:#fffdf7'><b>Chấm tự động bằng AI</b> "
        "<span class='muted'>— đã chấm · điểm AI 78/100</span> "
        "<span class='btn btn-org btn-sm' style='float:right'>Chấm lại</span></div>"
        "<div class='card'><b>Tài liệu đã nộp (11)</b> <span class='btn btn-green btn-sm' style='float:right'>⬇ Tải tất cả (ZIP)</span></div>"
        "<div class='card'><b>Phần C — Giảng dạy và hỗ trợ học tập</b> <span style='float:right;color:#c2410c;font-weight:bold'>31/40</span>"
        + tbl(["Tiêu chí", "Lượt AI", "Điểm AI", "Nhận xét AI", "Điều chỉnh của Hội đồng"], [
            ["C1 /10", "8 · 8", "<b>8</b>", "Giáo án rõ mục tiêu...", "<span class='in' style='display:inline-block;width:60px'>8</span> <span class='btn btn-slate btn-sm'>Lưu</span>"],
            ["C2 /10", "7.5 · 8", "<b>8</b>", "Slide đạt yêu cầu...", "<span class='in' style='display:inline-block;width:60px'>8</span> <span class='btn btn-slate btn-sm'>Lưu</span>"],
        ]) + "</div>"
        "<div class='center'><span class='btn btn-green'>PHÊ DUYỆT KẾT QUẢ</span></div>"))

    # 7. Hội đồng/Admin — bảng điểm
    s["scores"] = render("scores", screen("council",
        "<div class='h1'>Bảng điểm giảng viên <span class='muted' style='font-size:14px'>Đã chấm: 3/4 hồ sơ</span> "
        "<span class='btn btn-green btn-sm' style='float:right'>⬇ Tải Excel</span></div>"
        + tbl(["Mã GV", "Họ tên", "Đơn vị", "B", "C", "D", "E", "F", "G", "Tổng", "Mức", "Tính chất"], [
            ["GV003", "Lê Minh Châu", "CNTT", "18", "37", "9", "18", "4", "4", "<b style='color:#c2410c'>88</b>", "Dẫn dắt", "<span style='color:#b45309'>Tạm tính</span>"],
            ["GV001", "Nguyễn Văn An", "CNTT", "16", "31", "8", "16", "4", "3", "<b style='color:#c2410c'>78</b>", "Thành thạo", "<span style='color:#047857'>Chính thức</span>"],
            ["GV002", "Trần Thị Bình", "CNTT", "13", "26", "7", "13", "3", "3", "<b style='color:#c2410c'>65</b>", "Cơ bản", "<span style='color:#b45309'>Tạm tính</span>"],
            ["GV004", "Phạm Quốc Dũng", "QTKD", "–", "–", "–", "–", "–", "–", "–", "", "<span class='muted'>Chưa chấm</span>"],
        ])))

    # 8. Admin — vận hành
    s["admin_index"] = render("admin_index", screen("admin",
        "<div class='h1'>Bảng điều khiển vận hành</div>"
        + stats([("Giảng viên", "500"), ("Hồ sơ", "486"), ("Đã chấm", "120"), ("Phản hồi chờ", "2")])
        + "<div class='row' style='margin-top:14px'>"
        "<div class='col'><div class='card'><b>Quy trình chấm</b>"
        "<div style='font-size:13px;margin-top:8px'>1. Khóa hồ sơ &amp; kiểm tra hợp lệ <span class='btn btn-slate btn-sm'>Khóa ngay</span><br><br>"
        "2. Chấm AI lần lượt từng hồ sơ tại trang Thẩm định (kiểm soát tiến độ &amp; chi phí)<br><br>"
        "3. Hội đồng điều chỉnh &amp; phê duyệt<br><br>"
        "4. Công bố kết quả <span class='btn btn-green btn-sm'>Công bố</span></div></div></div>"
        "<div class='col'><div class='card'><b>Công cụ</b>"
        "<div style='font-size:13px;line-height:2'>📧 Gửi email nhắc hạn<br>👥 Quản lý giảng viên (import CSV)<br>"
        "🔑 Cấu hình AI, mốc thời gian &amp; rubric<br>💰 Thống kê sử dụng AI (token, chi phí)<br>"
        "📦 Tải sản phẩm giảng viên (ZIP)<br>📊 Dashboard &amp; xuất báo cáo<br>💾 Sao lưu dữ liệu</div></div></div></div>"))

    # 9. Admin — cấu hình AI
    s["admin_ai"] = render("admin_ai", screen("admin",
        "<div class='h1'>Cấu hình</div>"
        "<div class='card' style='border:2px solid #fed7aa'><b>🔑 Cấu hình AI chấm điểm</b>"
        + tbl(["", ""], [
            ["Bộ chấm hiện dùng", "<b style='color:#047857'>Claude API (chấm thật)</b>"],
            ["Model", "claude-opus-4-8"],
            ["API key", "<span style='color:#047857'>Đã cấu hình ••••a1b2 (nguồn: DB)</span>"],
        ]) +
        "<div style='margin-top:10px'><div class='lbl'>Nạp/đổi Claude API key</div>"
        "<div class='in'>sk-ant-... (để trống = giữ key hiện tại)</div>"
        "<div class='btn btn-org btn-sm' style='margin-top:8px'>Lưu cấu hình AI</div> "
        "<div class='btn btn-blue btn-sm'>⚡ Kiểm tra kết nối Claude</div></div>"
        "<div style='margin-top:12px;font-size:13px'>Sử dụng AI — Số lần gọi: <b>1.560</b> · Token: <b>34,2 triệu</b> · "
        "Ước tính chi phí: <b style='color:#c2410c'>$240</b> <span class='muted'>(~6.120.000đ)</span></div></div>"))

    # 10. Admin — người dùng / import
    s["admin_users"] = render("admin_users", screen("admin",
        "<div class='h1'>Người dùng (500)</div>"
        "<div class='row'><div class='col' style='flex:2'><div class='card' style='padding:0'>"
        + tbl(["Mã GV", "Họ tên", "Email", "Đơn vị", "Chức vụ", "Vai trò", ""], [
            ["GV001", "Nguyễn Văn An", "gv001@dainam.edu.vn", "CNTT", "Giảng viên", "Giảng viên", "Đặt mật khẩu · Khóa"],
            ["HD001", "Phạm Hội Đồng", "hoidong@dainam.edu.vn", "—", "Thành viên HĐ", "Hội đồng", "Đặt mật khẩu · Khóa"],
        ]) + "</div></div>"
        "<div class='col'><div class='card'><b>Import danh sách (CSV)</b>"
        "<div class='muted' style='font-size:12px;margin-top:4px'>Cột: ma_gv, ho_ten, email, don_vi, bo_mon, chuc_vu, role, password</div>"
        "<div class='in'>Chọn tệp .csv</div><div class='btn btn-org btn-sm' style='margin-top:8px'>Import</div></div></div></div>"))

    # 11. Admin — tải sản phẩm
    s["admin_dl"] = render("admin_dl", screen("admin",
        "<div class='h1'>Tải sản phẩm giảng viên "
        "<span class='btn btn-green btn-sm' style='float:right'>⬇ Tải toàn bộ tất cả giảng viên (ZIP)</span></div>"
        + tbl(["Mã GV", "Họ tên", "Đơn vị", "Tệp", "Liên kết", "Trạng thái", "Tải về"], [
            ["GV001", "Nguyễn Văn An", "CNTT", "11", "2", "Đã chấm", "<b style='color:#c2410c'>⬇ ZIP</b>"],
            ["GV002", "Trần Thị Bình", "CNTT", "9", "1", "Đã nộp", "<b style='color:#c2410c'>⬇ ZIP</b>"],
            ["GV003", "Lê Minh Châu", "CNTT", "12", "3", "Đã duyệt", "<b style='color:#c2410c'>⬇ ZIP</b>"],
        ])))

    return s


# ---------- dựng tài liệu Word ----------

def build(out_path: str) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    imgs = build_screens()
    today = datetime.now(ZoneInfo("Asia/Ho_Chi_Minh"))
    ORANGE = RGBColor(0xC2, 0x41, 0x0C)
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def h(text, size=14, color=ORANGE, before=10, after=4, align=None, bold=True):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    def step(n, text):
        p = doc.add_paragraph()
        p.add_run(f"Bước {n}. ").bold = True
        p.add_run(text)
        p.paragraph_format.space_after = Pt(3)

    def img(key, caption):
        doc.add_picture(str(imgs[key]), width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph()
        r = c.add_run(f"Hình: {caption}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)

    # Trang bìa
    h("HƯỚNG DẪN SỬ DỤNG PHẦN MỀM", size=20, align=WD_ALIGN_PARAGRAPH.CENTER, before=0)
    h("Hệ thống đánh giá năng lực ứng dụng AI của giảng viên — DNU AI-Assess",
      size=13, color=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(f"Trường Đại học Đại Nam · {today:%d/%m/%Y}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    doc.add_paragraph()

    h("Giới thiệu chung", size=14)
    doc.add_paragraph(
        "DNU AI-Assess là phần mềm tiếp nhận hồ sơ trực tuyến và chấm tự động bằng AI theo rubric của đề bài, "
        "phục vụ ba nhóm người dùng: Giảng viên (nộp bài), Hội đồng (chấm & phê duyệt) và Quản trị viên (vận hành). "
        "Tài liệu này hướng dẫn từng bước cho mỗi nhóm.")

    h("Đăng nhập (chung cho mọi vai trò)", size=14)
    step(1, "Truy cập địa chỉ hệ thống do Nhà trường cung cấp.")
    step(2, "Nhập ID đăng nhập (email hoặc mã giảng viên) và mật khẩu, bấm Đăng nhập. "
            "Hệ thống tự đưa tới trang phù hợp với vai trò.")
    step(3, "Lần đầu nên vào mục Đổi mật khẩu (góc trên bên phải) để đặt mật khẩu riêng.")
    img("login", "Màn hình đăng nhập bằng ID + mật khẩu.")

    # I. Giảng viên
    h("PHẦN I — HƯỚNG DẪN CHO GIẢNG VIÊN", size=15)
    step(1, "Vào Hồ sơ của tôi: xem tiến độ các Phần A–G và hạn nộp (17h00 ngày 30/6/2026).")
    step(2, "Bấm vào Phần A để kê khai thông tin (họ tên, mã GV, khoa/bộ môn, học phần, công cụ AI, mức tự đánh giá).")
    img("lec_dashboard", "Trang hồ sơ: tiến độ từng phần và nút NỘP HỒ SƠ.")
    step(3, "Mở từng Phần B–G: tải tệp sản phẩm (chọn được nhiều tệp một lần) và tải/ dán liên kết minh chứng sử dụng AI. "
            "Đặt tên tệp theo mẫu MãGV_TênPhần_TênSảnPhẩm; mỗi tệp tối đa 200MB.")
    img("lec_part", "Trang nộp một phần: khu vực sản phẩm và minh chứng AI.")
    step(4, "Kiểm tra đủ thành phần (báo 'Đủ thành phần' ở mỗi phần), sau đó bấm NỘP HỒ SƠ. "
            "Có thể nộp lại nhiều lần trước hạn — hệ thống chỉ chấm bản cuối.")
    step(5, "Sau khi Nhà trường công bố, vào Kết quả để xem điểm từng tiêu chí và nhận xét; "
            "có thể gửi phản hồi trong 03 ngày làm việc.")
    img("lec_result", "Trang kết quả: tổng điểm, mức năng lực, nhận xét từng tiêu chí.")

    # II. Hội đồng
    h("PHẦN II — HƯỚNG DẪN CHO HỘI ĐỒNG", size=15)
    step(1, "Vào Thẩm định: danh sách hồ sơ, lọc theo khoa/trạng thái. "
            "Hồ sơ ≥85 điểm hoặc có dấu hiệu bất thường được đánh dấu 'Thẩm định bắt buộc'.")
    img("council_list", "Danh sách thẩm định với điểm AI và cờ bắt buộc.")
    step(2, "Mở một hồ sơ, bấm Chấm tự động hồ sơ này để AI chấm theo rubric (mỗi tiêu chí 2 lượt, "
            "lệch >15% chấm lượt 3 lấy trung vị). Có thể tải toàn bộ tài liệu của giảng viên dạng ZIP.")
    step(3, "Xem điểm AI và nhận xét từng tiêu chí; nhập điểm điều chỉnh kèm lý do (nếu cần) rồi bấm Lưu.")
    step(4, "Khi hoàn tất, bấm PHÊ DUYỆT KẾT QUẢ. Điểm cuối là điểm Hội đồng phê duyệt; mọi thao tác được ghi vết.")
    img("council_detail", "Trang chấm chi tiết: điểm AI, điều chỉnh của Hội đồng, phê duyệt.")
    step(5, "Vào Bảng điểm để xem/đối chiếu điểm tất cả giảng viên (kể cả khi chưa chốt) và tải Excel. "
            "Mục Phản hồi để trả lời ý kiến của giảng viên.")
    img("scores", "Bảng điểm tất cả giảng viên — xem được trước khi chốt, tải Excel.")

    # III. Quản trị
    h("PHẦN III — HƯỚNG DẪN CHO QUẢN TRỊ VIÊN", size=15)
    step(1, "Vào Người dùng → Import danh sách (CSV) theo mẫu để tạo tài khoản giảng viên/hội đồng hàng loạt; "
            "có thể đặt lại mật khẩu, khóa/mở tài khoản.")
    img("admin_users", "Quản lý người dùng và import danh sách từ CSV.")
    step(2, "Vào Cấu hình AI → nhập Claude API key của Trường, chọn model, bấm Lưu rồi Kiểm tra kết nối. "
            "App sẽ dùng key này để chấm.")
    img("admin_ai", "Nạp API key AI và xem nhanh chi phí đã dùng.")
    step(3, "Theo dõi Thống kê sử dụng AI (số lần gọi, token, ước tính chi phí USD/VND) để kiểm soát ngân sách.")
    step(4, "Trang Vận hành: Khóa hồ sơ khi đến hạn; điều phối quá trình chấm (Hội đồng/Admin chấm lần lượt từng bài); "
            "Công bố kết quả để gửi email cho giảng viên.")
    img("admin_index", "Trang vận hành: quy trình chấm và bộ công cụ quản trị.")
    step(5, "Trang Tải sản phẩm: tải toàn bộ bài của từng giảng viên hoặc tất cả (lọc theo khoa) dưới dạng ZIP. "
            "Trang Báo cáo/Bảng điểm để xem dashboard, phân loại 4 mức và xuất Excel toàn trường.")
    img("admin_dl", "Tải sản phẩm giảng viên theo từng người hoặc toàn bộ.")

    # Ghi chú
    h("Ghi chú", size=13)
    note = doc.add_paragraph(
        "Các hình trong tài liệu là ảnh minh họa giao diện, có thể thay bằng ảnh chụp màn hình thực tế trên "
        "hệ thống đã triển khai của Nhà trường. Chức năng và nhãn thao tác đúng với phần mềm bàn giao.")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(out_path)
    print("Đã tạo:", out_path)


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "DNU-AI-Assess-HuongDanSuDung.docx")
