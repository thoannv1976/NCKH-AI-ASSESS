"""Chuyển tài liệu Markdown sang Word (.docx) — hỗ trợ tiêu đề, đoạn, danh sách,
khối mã, bảng, in đậm và mã inline. Dùng để xuất tài liệu bàn giao.

Chạy: python scripts/md_to_docx.py <input.md> <output.docx> ["Tiêu đề bìa"]
"""
from __future__ import annotations

import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ORANGE = RGBColor(0xC2, 0x41, 0x0C)
CODE_BG = "F1F5F9"


def _add_runs(p, text: str) -> None:
    """Thêm text vào đoạn, xử lý **đậm** và `mã inline`."""
    for part in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        else:
            p.add_run(part)


def _shade(cell_or_para, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pr = cell_or_para
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pr.append(shd)


def convert(md_path: str, out_path: str, cover_title: str | None = None) -> None:
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    if cover_title:
        p = doc.add_paragraph()
        r = p.add_run(cover_title)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = ORANGE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Khối mã ```
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(6)
            _shade(p._p.get_or_add_pPr(), CODE_BG)
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            continue

        # Bảng | ... |
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            for ci, htext in enumerate(header):
                cell = t.rows[0].cells[ci]
                cell.text = ""
                _add_runs(cell.paragraphs[0], htext)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)
            for row in rows:
                cells = t.add_row().cells
                for ci, val in enumerate(row):
                    if ci < len(cells):
                        cells[ci].text = ""
                        _add_runs(cells[ci].paragraphs[0], val)
                        for run in cells[ci].paragraphs[0].runs:
                            run.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # Tiêu đề
        if stripped.startswith("### "):
            h = doc.add_paragraph()
            r = h.add_run(stripped[4:])
            r.bold = True
            r.font.size = Pt(12)
        elif stripped.startswith("## "):
            h = doc.add_paragraph()
            r = h.add_run(stripped[3:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = ORANGE
            h.paragraph_format.space_before = Pt(8)
        elif stripped.startswith("# "):
            h = doc.add_paragraph()
            r = h.add_run(stripped[2:])
            r.bold = True
            r.font.size = Pt(17)
            r.font.color.rgb = ORANGE
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Danh sách số
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
        # Danh sách chấm
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
        # Đường kẻ / trống
        elif stripped in ("", "---", "***"):
            if stripped == "":
                pass
        # Đoạn thường (gộp dòng liền kề)
        else:
            para = stripped
            while i + 1 < n and lines[i + 1].strip() and not re.match(
                r"^(#|\d+\.\s|[-*]\s|\||```)", lines[i + 1].strip()):
                i += 1
                para += " " + lines[i].strip()
            italic = para.startswith("*") and para.endswith("*") and not para.startswith("**")
            p = doc.add_paragraph()
            _add_runs(p, para.strip("*") if italic else para)
            if italic:
                for r in p.runs:
                    r.italic = True
                    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    r.font.size = Pt(9)
        i += 1

    doc.save(out_path)
    print("Đã tạo:", out_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Cách dùng: python scripts/md_to_docx.py <input.md> <output.docx> [tiêu đề bìa]")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
