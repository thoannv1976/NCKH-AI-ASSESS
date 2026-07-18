"""Tạo dữ liệu demo: người dùng 3 vai trò + 1 hồ sơ mẫu đầy đủ (kèm tệp docx thật)
để demo trọn luồng nộp → khóa → chấm → thẩm định → công bố.

Chạy:  python scripts/seed_demo.py
"""
from __future__ import annotations

import io
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.config import GRADED_PARTS, get_settings, now_vn  # noqa: E402
from app.db import create_store, new_id  # noqa: E402
from app.rubric import get_rubric  # noqa: E402
from app.storage import create_storage  # noqa: E402

USERS = [
    {"ma_gv": "GV001", "ho_ten": "Nguyễn Văn An", "email": "gv001@dainam.edu.vn",
     "khoa": "Công nghệ thông tin", "bo_mon": "Kỹ thuật phần mềm", "chuc_vu": "Giảng viên", "role": "lecturer"},
    {"ma_gv": "GV002", "ho_ten": "Trần Thị Bình", "email": "gv002@dainam.edu.vn",
     "khoa": "Công nghệ thông tin", "bo_mon": "Hệ thống thông tin", "chuc_vu": "Trưởng bộ môn", "role": "lecturer"},
    {"ma_gv": "GV003", "ho_ten": "Lê Minh Châu", "email": "gv003@dainam.edu.vn",
     "khoa": "Quản trị kinh doanh", "bo_mon": "Marketing", "chuc_vu": "Giảng viên", "role": "lecturer"},
    {"ma_gv": "HD001", "ho_ten": "Phạm Hội Đồng", "email": "hoidong@dainam.edu.vn",
     "khoa": "", "bo_mon": "", "chuc_vu": "Thành viên Hội đồng", "role": "council"},
    {"ma_gv": "QT001", "ho_ten": "Quản Trị Viên", "email": "admin@dainam.edu.vn",
     "khoa": "", "bo_mon": "", "chuc_vu": "Quản trị hệ thống", "role": "admin"},
]

SAMPLE_DOCS = {
    "B": ("DeCuongHocPhan", [
        "ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN: NHẬP MÔN TRÍ TUỆ NHÂN TẠO (CẬP NHẬT 2026-2027)",
        "CLO1: Trình bày được các khái niệm nền tảng của AI (tương thích PLO2).",
        "CLO2: Vận dụng được thuật toán tìm kiếm và học máy cơ bản vào bài toán thực tế (PLO3, PLO5).",
        "CLO3: Đánh giá được tác động đạo đức của hệ thống AI (PLO7).",
        "Ma trận CLO–PLO: CLO1-PLO2(H); CLO2-PLO3(H),PLO5(M); CLO3-PLO7(M).",
        "Kế hoạch 15 tuần: Tuần 1 Tổng quan AI; Tuần 2-4 Tìm kiếm; Tuần 5-9 Học máy; Tuần 10-12 Học sâu; "
        "Tuần 13-14 Đạo đức AI; Tuần 15 Tổng kết dự án.",
        "Bảng so sánh trước–sau: bổ sung chương Generative AI (AI đề xuất, đã chấp nhận có điều chỉnh); "
        "cập nhật học liệu 2024-2026 (AI đề xuất); giữ nguyên chuẩn đánh giá quá trình (từ chối đề xuất AI vì trùng quy chế).",
        "Báo cáo đối sánh: so với chương trình CS50 AI (Harvard) và Khoa học máy tính (ĐH Bách khoa HN). "
        "Đề xuất 3 cải tiến: (1) thêm dự án capstone với doanh nghiệp; (2) tích hợp công cụ AI hỗ trợ lập trình; "
        "(3) chuẩn đầu ra về prompt engineering.",
    ]),
    "C": ("KeHoachBaiGiang", [
        "KẾ HOẠCH BÀI GIẢNG: THUẬT TOÁN TÌM KIẾM A* (Buổi 6, 3 tiết)",
        "Mục tiêu: SV cài đặt được A*, so sánh heuristic (CLO2).",
        "Hoạt động 1 (Khởi động, 15p): SV dùng chatbot lớp học đoán đường đi tối ưu trên bản đồ — AI sinh tình huống.",
        "Hoạt động 2 (Thực hành, 60p): SV dùng trợ lý AI debug cài đặt A*; giảng viên cung cấp bộ prompt mẫu.",
        "Slide 18 trang kèm sơ đồ tư duy do AI hỗ trợ tạo (đính kèm).",
        "Trợ lý AI 'AI-Tutor DNU': bộ 12 prompt hướng dẫn tự học kèm hướng dẫn sử dụng; đã thử nghiệm với 5 SV.",
        "Phương án phân hóa: nhóm khá-giỏi nhận bài toán mở rộng (AI sinh test nâng cao, phản hồi sâu); "
        "nhóm cần hỗ trợ học theo lộ trình chậm với gợi ý từng bước từ chatbot; AI điều chỉnh nhịp độ qua quiz chẩn đoán.",
    ]),
    "D": ("NganHangCauHoi", [
        "NGÂN HÀNG CÂU HỎI CHƯƠNG 3: HỌC MÁY CƠ BẢN",
        "32 câu trắc nghiệm: 12 câu mức Nhớ, 10 câu mức Hiểu, 7 câu Vận dụng, 3 câu Phân tích (danh sách kèm đáp án).",
        "5 câu tự luận/tình huống: phân tích bài toán phân loại email rác; thiết kế pipeline dữ liệu...",
        "Ghi chú hiệu chỉnh: 8 câu AI tạo bị sửa đáp án nhiễu; 3 câu bị loại vì ngoài phạm vi.",
        "Rubric chấm bài tập lớn: 4 tiêu chí x 4 mức kèm mô tả; 2 bài làm mẫu do AI mô phỏng (giỏi/trung bình) "
        "và kết quả chấm thử: 8.5/10 và 6.0/10, nhất quán rubric.",
        "3 biện pháp kiểm soát lạm dụng AI: vấn đáp ngẫu nhiên 10%; yêu cầu nộp nhật ký prompt; "
        "đề bài gắn dữ liệu riêng từng nhóm.",
    ]),
    "E": ("TongQuanTaiLieu", [
        "TỔNG QUAN TÀI LIỆU: ỨNG DỤNG LLM TRONG GIÁO DỤC ĐẠI HỌC (2021-2026)",
        "Bảng tổng hợp 16 tài liệu (15/16 trong 5 năm gần nhất) gồm DOI đã kiểm chứng từng mục.",
        "Khoảng trống nghiên cứu: thiếu nghiên cứu thực nghiệm về AI hỗ trợ giảng viên khối kinh tế tại Việt Nam.",
        "Ghi chú kiểm chứng: phát hiện 2 trích dẫn AI tạo không tồn tại, đã loại và thay bằng nguồn thực (có DOI).",
        "ĐỀ CƯƠNG NGHIÊN CỨU: 'Tác động của trợ lý AI đến kết quả học tập SV DNU' — câu hỏi nghiên cứu, "
        "thiết kế bán thực nghiệm 2 nhóm, dữ liệu khảo sát + điểm số, kế hoạch 8 tháng.",
        "Nhiệm vụ khuyến khích: dùng AI phân tích tập dữ liệu khảo sát mẫu 120 phản hồi (kèm script và biểu đồ).",
    ]),
    "F": ("CamNangCongDong", [
        "CẨM NANG: ỨNG DỤNG AI AN TOÀN CHO DOANH NGHIỆP NHỎ (sản phẩm phục vụ cộng đồng)",
        "12 trang hướng dẫn chọn công cụ AI, bảo mật dữ liệu, mẫu prompt cho 5 nghiệp vụ phổ biến.",
        "Đối tượng: doanh nghiệp nhỏ tại Hà Đông; đã chia sẻ qua kênh kết nối doanh nghiệp của DNU.",
    ]),
    "G": ("BaoCaoKinhNghiem", [
        "BÁO CÁO BÀI HỌC KINH NGHIỆM SỬ DỤNG AI",
        "Công cụ đã dùng: Claude (soạn đề cương, rubric), ChatGPT (sinh câu hỏi), Gamma (slide), NotebookLM (tổng quan).",
        "AI làm tốt: dàn ý, ma trận, mô phỏng bài làm. AI chưa tốt: trích dẫn khoa học (2 trích dẫn ảo), "
        "ngữ cảnh chương trình DNU phải cung cấp thủ công.",
        "Kỹ thuật prompt hiệu quả nhất: cung cấp vai trò + ngữ cảnh + ví dụ mẫu + yêu cầu phản biện lại kết quả.",
        "Cam kết: kiểm chứng mọi nội dung AI trước khi dùng; minh bạch với sinh viên về việc dùng AI.",
        "DANH MỤC MINH CHỨNG TỔNG HỢP: Phần B (2 liên kết hội thoại); Phần C (4 nhật ký prompt + ảnh thử nghiệm chatbot); "
        "Phần D (nhật ký prompt); Phần E (nhật ký + danh sách DOI kiểm chứng); Phần F (nhật ký prompt).",
    ]),
}


def make_docx(title: str, paragraphs: list[str]) -> bytes:
    import docx

    d = docx.Document()
    d.add_heading(title, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def seed(store, storage) -> None:
    """Tạo người dùng demo + hồ sơ mẫu trên store/storage cho trước (idempotent)."""
    get_rubric(store)

    from app.security import hash_password

    demo_pw = hash_password("demo123")  # mật khẩu demo dùng chung cho mọi tài khoản mẫu
    users_by_email = {}
    for u in USERS:
        existing = store.find_one("users", email=u["email"])
        if existing:
            # Backfill mật khẩu nếu tài khoản cũ chưa có (vd dữ liệu từ lần deploy trước)
            if not existing.get("password_hash"):
                store.patch("users", existing["id"], {"password_hash": demo_pw, "active": True})
                print(f"~ backfill mật khẩu demo123 cho {u['email']}")
            users_by_email[u["email"]] = store.get("users", existing["id"])
            continue
        uid = store.add("users", {**u, "active": True, "password_hash": demo_pw})
        users_by_email[u["email"]] = store.get("users", uid)
        print(f"+ user {u['email']} ({u['role']}) — mật khẩu: demo123")

    # Hồ sơ mẫu đầy đủ cho GV001
    owner = users_by_email["gv001@dainam.edu.vn"]
    if store.find_one("submissions", user_id=owner["id"]):
        print("Hồ sơ mẫu GV001 đã tồn tại — bỏ qua.")
        return
    sid = new_id()
    store.put("submissions", sid, {
        "id": sid, "user_id": owner["id"], "status": "submitted",
        "part_a": {
            "ho_ten": owner.get("ho_ten", ""), "ma_gv": owner.get("ma_gv", ""),
            "khoa_bo_mon": f"{owner.get('khoa', '')} - {owner.get('bo_mon', '')}",
            "hoc_phan": "Nhập môn Trí tuệ nhân tạo (HK1 2026-2027)",
            "cong_cu_ai": ["Claude", "ChatGPT", "Gamma", "NotebookLM"],
            "muc_thanh_thao": 4,
        },
        "created_at": now_vn().isoformat(), "submitted_at": now_vn().isoformat(),
    })
    for part in GRADED_PARTS:
        name, paras = SAMPLE_DOCS[part]
        fname = f"GV001_Phan{part}_{name}.docx"
        key = f"{sid}/{part}/product/{int(time.time())}_{fname}"
        storage.save(key, io.BytesIO(make_docx(f"Phần {part}: {name}", paras)))
        store.add("submission_items", {
            "submission_id": sid, "part": part, "kind": "product", "type": "file",
            "storage_path": key, "original_name": fname,
            "size": 20000, "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "uploaded_at": now_vn().isoformat(),
        })
        if part != "G":
            ev_name = f"GV001_Phan{part}_NhatKyPrompt.docx"
            ev_key = f"{sid}/{part}/evidence/{int(time.time())}_{ev_name}"
            storage.save(ev_key, io.BytesIO(make_docx(
                f"Nhật ký prompt Phần {part}",
                [f"Công cụ: Claude (phiên bản web, 6/2026). Chuỗi {3 + ord(part) % 3} prompt chính:",
                 "1) Cung cấp ngữ cảnh học phần và yêu cầu nháp ban đầu.",
                 "2) Phản biện kết quả, yêu cầu đối chiếu chuẩn đầu ra.",
                 "3) Tinh chỉnh theo góp ý của giảng viên, chốt phiên bản cuối.",
                 "Liên kết hội thoại: https://claude.ai/share/vi-du-demo"]
            )))
            store.add("submission_items", {
                "submission_id": sid, "part": part, "kind": "evidence", "type": "file",
                "storage_path": ev_key, "original_name": ev_name,
                "size": 12000, "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "uploaded_at": now_vn().isoformat(),
            })
    print(f"+ hồ sơ mẫu GV001 ({sid}) với {len(GRADED_PARTS)} phần (sản phẩm + minh chứng docx)")


def main() -> None:
    settings = get_settings()
    seed(create_store(settings), create_storage(settings))
    print("\nXong. Chạy: uvicorn app.main:app --reload  → http://localhost:8000 (đăng nhập tài khoản demo)")


if __name__ == "__main__":
    main()
