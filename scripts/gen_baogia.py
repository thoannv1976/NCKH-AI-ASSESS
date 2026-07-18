"""Sinh bản ĐỀ XUẤT CHUYỂN GIAO / BÁO GIÁ phần mềm DNU AI-Assess ra file .docx.

Chạy: python scripts/gen_baogia.py [đường_dẫn_xuất.docx]
Số liệu khớp docs/CHI_PHI_CHUYEN_GIAO.md.
"""
from __future__ import annotations

import sys
from datetime import datetime
from zoneinfo import ZoneInfo

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ORANGE = RGBColor(0xEA, 0x58, 0x0C)


def _h(doc, text, size=13, color=None, align=None, bold=True, space_before=8, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _table(doc, headers, rows, widths=None, total_row=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htext
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        is_last = total_row and ri == len(rows) - 1
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                if i > 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
                    if is_last:
                        run.bold = True
    return t


def build(path: str) -> None:
    today = datetime.now(ZoneInfo("Asia/Ho_Chi_Minh"))
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Tiêu đề
    _h(doc, "ĐỀ XUẤT CHUYỂN GIAO PHẦN MỀM", size=17, color=ORANGE,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
    _h(doc, "Hệ thống đánh giá năng lực ứng dụng AI của giảng viên — DNU AI-Assess",
       size=12, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    p = doc.add_paragraph(f"Ngày {today:%d} tháng {today:%m} năm {today:%Y}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True

    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: BAN GIÁM HIỆU TRƯỜNG ĐẠI HỌC ĐẠI NAM").runs[0].bold = True
    intro = doc.add_paragraph(
        "Trên cơ sở yêu cầu xây dựng hệ thống tiếp nhận và chấm tự động bằng AI cho Chương trình "
        "khảo sát, đánh giá năng lực ứng dụng AI của giảng viên năm 2026, chúng tôi trân trọng gửi tới "
        "Quý Trường đề xuất chuyển giao phần mềm DNU AI-Assess với nội dung và báo giá như sau."
    )
    intro.paragraph_format.space_after = Pt(6)

    # I. Giới thiệu
    _h(doc, "I. GIỚI THIỆU HỆ THỐNG", size=13, color=ORANGE)
    doc.add_paragraph(
        "DNU AI-Assess là ứng dụng web phục vụ trọn vòng đời đợt đánh giá: tiếp nhận hồ sơ trực tuyến "
        "theo cấu trúc Phần A–G, chấm tự động bằng Claude AI theo rubric của đề bài, hỗ trợ Hội đồng "
        "thẩm định, tổng hợp kết quả và xuất báo cáo.")
    for line in [
        "Nộp hồ sơ A–G: kê khai, tải nhiều tệp/liên kết, kiểm tra hợp lệ thời gian thực, nộp lại trước hạn.",
        "Chấm AI theo rubric 4 mức: mỗi tiêu chí 2 lượt độc lập, lệch >15% chấm lượt 3 lấy trung vị; "
        "trừ điểm khi thiếu minh chứng; phát hiện dấu hiệu bất thường.",
        "Admin/Hội đồng chấm lần lượt từng hồ sơ để kiểm soát tiến độ và chi phí; điều chỉnh điểm, phê duyệt, ghi vết.",
        "Báo cáo: bảng điều khiển theo khoa, phân loại 4 mức năng lực, xuất Excel, hồ sơ năng lực, tải sản phẩm (ZIP).",
        "Quản trị: nạp API key AI trong app, thống kê token & ước tính chi phí, quản lý người dùng, cấu hình mốc thời gian.",
        "Đăng nhập bằng ID + mật khẩu, phân quyền 3 vai trò; ~56 kiểm thử tự động; bàn giao kèm mã nguồn.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # II. Phạm vi chuyển giao
    _h(doc, "II. PHẠM VI CHUYỂN GIAO", size=13, color=ORANGE)
    for line in [
        "Toàn bộ mã nguồn ứng dụng + quyền sử dụng vĩnh viễn trong nội bộ Trường Đại học Đại Nam.",
        "Triển khai, cấu hình hệ thống trên hạ tầng Google Cloud của Trường (Cloud Run, Firestore, "
        "Cloud Storage, Cloud Scheduler); nạp API key AI; import danh sách giảng viên.",
        "Tài liệu hướng dẫn vận hành; tập huấn cho quản trị viên, Hội đồng và đầu mối hỗ trợ giảng viên.",
        "Bảo hành 6 tháng kể từ ngày nghiệm thu (sửa lỗi, hỗ trợ kỹ thuật trong đợt đánh giá đầu).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # III. Báo giá gói chuyển giao
    _h(doc, "III. BÁO GIÁ GÓI CHUYỂN GIAO (một lần)", size=13, color=ORANGE)
    _table(doc, ["Hạng mục", "Thành tiền (VNĐ)"], [
        ["1. Bản quyền + chuyển giao mã nguồn (vĩnh viễn, nội bộ DNU)", "85.000.000"],
        ["2. Triển khai & cấu hình trên hạ tầng GCP của Trường", "15.000.000"],
        ["3. Đào tạo, tài liệu & bảo hành 6 tháng", "10.000.000"],
        ["CỘNG GÓI CHUYỂN GIAO (chưa gồm VAT)", "110.000.000"],
    ], total_row=True)
    doc.add_paragraph("Bằng chữ: Một trăm mười triệu đồng chẵn (chưa bao gồm VAT).").runs[0].italic = True

    # IV. Chi phí vận hành DNU chi trả
    _h(doc, "IV. CHI PHÍ VẬN HÀNH (Trường chi trả trực tiếp)", size=13, color=ORANGE)
    doc.add_paragraph(
        "Phần mềm cho phép Trường tự nạp API key AI và chạy trên tài khoản Google Cloud của Trường, "
        "do đó chi phí AI và hạ tầng được tính trực tiếp vào tài khoản của Trường (không qua bên cung cấp). "
        "Dự toán cho một đợt đánh giá 500 giảng viên (số liệu đo thực tế từ hệ thống):")
    _table(doc, ["Khoản mục", "Dự toán (VNĐ)"], [
        ["Chấm AI 500 hồ sơ (≈ 2 USD/hồ sơ × 500) + dự phòng 20%", "~30.000.000"],
        ["Google Cloud (Cloud Run, Firestore, Storage — ~2 tháng cao điểm)", "~2.500.000 – 5.000.000"],
        ["CỘNG VẬN HÀNH / ĐỢT (chưa tối ưu)", "~32.000.000 – 35.000.000"],
        ["Nếu bật Batches API + prompt caching (giảm 40–50% chi phí AI)", "~18.000.000 – 22.000.000"],
    ], total_row=False)
    doc.add_paragraph(
        "Tương đương 40.000–70.000đ/giảng viên/đợt, tùy độ dài tài liệu và số lần chấm lại.").runs[0].italic = True

    # V. Hạng mục nâng cấp tùy chọn
    _h(doc, "V. HẠNG MỤC NÂNG CẤP QUY MÔ LỚN (tùy chọn)", size=13, color=ORANGE)
    doc.add_paragraph(
        "Để vận hành ổn định ở giờ cao điểm 300–500 giảng viên nộp đồng thời, đề xuất bổ sung:")
    _table(doc, ["Hạng mục", "Thành tiền (VNĐ)"], [
        ["Tải tệp thẳng lên Cloud Storage (signed URL) — bỏ giới hạn 32MB", "12.000.000 – 20.000.000"],
        ["Hàng đợi chấm + Batches API + chạy song song (giảm 50% chi phí AI)", "20.000.000 – 35.000.000"],
        ["Đọc PDF/ảnh bằng AI (vision) + kiểm thử tải", "10.000.000 – 18.000.000"],
        ["Email dịch vụ + quên mật khẩu + sao lưu tự động", "8.000.000 – 15.000.000"],
        ["CỘNG GÓI NÂNG CẤP", "~50.000.000 – 88.000.000"],
    ], total_row=True)

    # VI. Phương án hợp tác
    _h(doc, "VI. CÁC PHƯƠNG ÁN HỢP TÁC", size=13, color=ORANGE)
    _table(doc, ["Phương án", "Hình thức", "Chi phí"], [
        ["A. Chuyển giao trọn gói (khuyến nghị)",
         "Trường sở hữu phần mềm + mã nguồn, tự vận hành", "~110 triệu một lần + AI/cloud Trường tự trả"],
        ["B. Dịch vụ theo đầu giảng viên",
         "Bên cung cấp vận hành trọn gói", "120.000–180.000đ/GV/đợt (500 GV ≈ 60–90 triệu/đợt)"],
        ["C. Thuê bao SaaS hằng năm",
         "Nền tảng + cập nhật + hỗ trợ", "~60–90 triệu/năm + AI thực dùng"],
    ])
    doc.add_paragraph(
        "Khuyến nghị Phương án A — phù hợp đơn vị giáo dục sử dụng lâu dài; minh bạch chi phí; "
        "Trường chủ động vận hành trên hạ tầng và API key của mình.").runs[0].italic = True

    # VII. Điều khoản
    _h(doc, "VII. ĐIỀU KHOẢN CHUNG", size=13, color=ORANGE)
    for line in [
        "Thời gian triển khai: 5–7 ngày làm việc kể từ khi ký hợp đồng và Trường cấp quyền hạ tầng GCP.",
        "Bảo hành: 6 tháng kể từ nghiệm thu. Bảo trì/hỗ trợ năm tiếp theo (tùy chọn): ~18 triệu/năm (≈18% giá license).",
        "Thanh toán: 50% khi ký hợp đồng, 50% khi nghiệm thu. Hình thức: chuyển khoản.",
        "Báo giá có hiệu lực 30 ngày; giá chưa bao gồm thuế VAT; chi phí AI và Google Cloud do Trường chi trả trực tiếp.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # Ký
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    left, right = t.rows[0].cells
    left.text = "ĐẠI DIỆN TRƯỜNG ĐẠI HỌC ĐẠI NAM\n(Ký, ghi rõ họ tên)"
    right.text = "ĐẠI DIỆN ĐƠN VỊ CUNG CẤP\n(Ký, ghi rõ họ tên)"
    for c in t.rows[0].cells:
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Ghi chú: Các con số là dự toán tham khảo, chốt theo hợp đồng và phạm vi cụ thể. "
        "Chi phí AI thay đổi theo độ dài tài liệu, số lần chấm lại và model lựa chọn.")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(path)
    print("Đã tạo:", path)


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "DNU-AI-Assess-DeXuatChuyenGiao.docx")
